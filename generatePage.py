import sys
import argparse

from docx import Document
from docx.shared import Inches

pars = argparse.ArgumentParser()
pars.add_argument("--start_year", help="year of birth", required=True)
pars.add_argument("--years_per_page",  required=False, help="e.g. 2 years a page")

document = Document()

document.add_picture('header.png', width=Inches(7.25))

def printTutorial():
    document.add_paragraph().add_run("This journal is meant to be a trip down memory lane. It\'s supposed to be great at sparking stories and conversations between people close to you. It can even document a life. You can fill it out at any pace. The writer can skip around, but it starts from most recent because those are the memories we remember more directly. Diagrams and pointers can be drawn on the margin for a more robust form of journaling. There are sections at the end to add standout, timeless stories or little lessons. ")
    printYear("Demo")
    r = document.add_paragraph().add_run("This March I learned how to play mahjong. I left position at ____ and started a\nnew type of job at ____. When I first started I remember this story where _____.\nMy daily routine consisted of _____. A cool TV series that I remember watching\naround this time was ______. I traveled around the city and saw ____ in summer...")
    r.underline = True

def printYear(year):
    document.add_heading(str(year), level=1)
    p = document.add_paragraph('These are the remarks and the thoughts about ')
    p.add_run(str(year)+". \t\t\t\t\tYou can add \t\t\t\t\t\t\t\t\t\t\tdrawings / diagrams here.").italic = True
    if year == 2020:
        document.add_paragraph().add_run("This year has not happened yet, so write what you look forward to!")
        
def printLine():
    p = document.add_paragraph()
    run = p.add_run('                                                                        \
                                                             ')
    run.underline = True
    

args = pars.parse_args()
start = args.start_year

printTutorial()
for year in range(2020,int(start)-1, -1):

    printYear(year)
    for line in range(64):
        printLine()
    if year != 2020: 
        document.add_paragraph().add_run("The part about this year I was most grateful for:").italic
        for line in range(16):
            printLine()
        if year % 3 == 0:
            document.add_picture('header.png', width=Inches(7.25))

for line in range(256):
        printLine()

margin=0.75
sections = document.sections
for section in sections:
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

document.save('demo.docx')